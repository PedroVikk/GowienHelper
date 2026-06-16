"""Testes dos extractors e do serviço de extração."""
import pytest

from app.core.exceptions import ValidationError
from app.infrastructure.extraction.extractors import (
    DocxExtractor,
    ImageOcrExtractor,
    PdfExtractor,
    PlainTextExtractor,
)
from app.infrastructure.extraction.service import ExtractionService


def test_plain_text_extractor(tmp_path):
    f = tmp_path / "nota.txt"
    f.write_text("Olá mundo\nlinha 2", encoding="utf-8")
    assert PlainTextExtractor().extract(str(f)) == "Olá mundo\nlinha 2"


def test_docx_extractor(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Resumo de Psicologia")
    doc.add_paragraph("Behaviorismo")
    path = tmp_path / "doc.docx"
    doc.save(str(path))

    text = DocxExtractor().extract(str(path))
    assert "Psicologia" in text and "Behaviorismo" in text


def test_service_dispatch_and_unsupported(tmp_path):
    f = tmp_path / "a.md"
    f.write_text("# Título", encoding="utf-8")
    svc = ExtractionService()
    assert svc.extract(str(f), "md") == "# Título"
    with pytest.raises(ValidationError):
        svc.extract(str(f), "video")


def test_can_handle_mapping():
    assert PdfExtractor().can_handle("pdf")
    assert ImageOcrExtractor().can_handle("image")
    assert PlainTextExtractor().can_handle("txt")
    assert PlainTextExtractor().can_handle("md")
    assert not PdfExtractor().can_handle("txt")
